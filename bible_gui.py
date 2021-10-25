import json
import os
import time
# from mttkinter import mtTkinter as tkinter
import tkinter
from tkinter import *
import threading
import tkinter.messagebox


from docx import Document
import re

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from cmd_colors import Color
from progressbar import *










class JwfwClient:
    def __init__(self):

        # 加载文档数据
        self.bible_data = Document("../data/bible.docx")
        self.bible_travel_data = Document("../data/bible_travel_handle.docx")
        self.filename = r'../data/dict.json'
        self.fn = r'../data/result_dict.json'

        # 验证章节是否有效的dict
        self.valid_dict = json.load(open(self.filename))
        self.row_dict = json.load(open(self.fn))

        # cmd颜色改变类
        self.clr = Color()

        # 卷名以及简写
        self.bible_dict = {"创世记": "创", "出埃及记": "出", "利未记": "利", "民数记": "民", "申命记": "申", "约书亚记": "书", "士师记": "士", "路得记": "得",
                      "撒母耳记上": "撒上", "撒母耳记下": "撒下", "列王记上": "王上",
                      "列王记下": "王下", "历代志上": "代上", "历代志下": "代下", "以斯拉记": "拉", "尼希米记": "尼", "以斯帖记": "斯", "约伯记": "伯",
                      "诗篇": "诗",
                      "箴言": "箴", "传道书": "传", "雅歌": "歌", "以赛亚书": "赛",
                      "耶利米书": "耶", "耶利米哀歌": "哀", "以西结书": "结", "但以理书": "但", "何西阿书": "何", "约珥书": "珥", "阿摩司书": "摩",
                      "俄巴底亚书": "俄",
                      "约拿书": "拿", "弥迦书": "弥", "那鸿书": "鸿", "哈巴谷书": "哈",
                      "西番雅书": "番", "哈该书": "该", "撒迦利亚书": "亚", "玛拉基书": "玛", "马太福音": "太", "马可福音": "可", "路加福音": "路",
                      "约翰福音": "约",
                      "使徒行传": "徒", "罗马书": "罗", "哥林多前书": "林前", "哥林多后书": "林后",
                      "加拉太书": "加", "以弗所书": "弗", "腓立比书": "腓", "歌罗西书": "西", "帖撒尼罗迦前书": "帖前", "帖撒尼罗迦后书": "帖后",
                      "提摩太前书": "提前",
                      "提摩太后书": "提后", "提多书": "多", "腓利门书": "门", "希伯来书": "来",
                      "雅各书": "雅", "彼得前书": "彼前", "彼得后书": "彼后", "约翰一书": "约壹", "约翰二书": "约贰", "约翰三书": "约叁", "犹大书": "犹",
                      "启示录": "启"}

        self.index_dict = {"1": "创世记", "2": "出埃及记", "3": "利未记", "4": "民数记", "5": "申命记", "6": "约书亚记", "7": "士师记", "8": "路得记",
                      "9": "撒母耳记上", "10": "撒母耳记下", "11": "列王记上",
                      "12": "列王记下", "13": "历代志上", "14": "历代志下", "15": "以斯拉记", "16": "尼希米记", "17": "以斯帖记", "18": "约伯记",
                      "19": "诗篇",
                      "20": "箴言", "21": "传道书", "22": "雅歌", "23": "以赛亚书",
                      "24": "耶利米书", "25": "耶利米哀歌", "26": "以西结书", "27": "但以理书", "28": "何西阿书", "29": "约珥书", "30": "阿摩司书",
                      "31": "俄巴底亚书",
                      "32": "约拿书", "33": "弥迦书", "34": "那鸿书", "35": "哈巴谷书",
                      "36": "西番雅书", "37": "哈该书", "38": "撒迦利亚书", "39": "玛拉基书", "40": "马太福音", "41": "马可福音", "42": "路加福音",
                      "43": "约翰福音",
                      "44": "使徒行传", "45": "罗马书", "46": "哥林多前书", "47": "哥林多后书",
                      "48": "加拉太书", "49": "以弗所书", "50": "腓立比书", "51": "歌罗西书", "52": "帖撒尼罗迦前书", "53": "帖撒尼罗迦后书",
                      "54": "提摩太前书",
                      "55": "提摩太后书", "56": "提多书", "57": "腓利门书", "58": "希伯来书",
                      "59": "雅各书", "60": "彼得前书", "61": "彼得后书", "62": "约翰一书", "63": "约翰二书", "64": "约翰三书", "65": "犹大书",
                      "66": "启示录"}

        self.number_list = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]

        self.num_dict = {"0": "", "1": "一", "2": "二", "3": "三", "4": "四", "5": "五", "6": "六", "7": "七", "8": "八", "9": "九",
                    "10": "十"}
        # paragraph = document.add_paragraph(bible_dict[bible_name])
        self.current_volume_name = ""
        self.current_verse = ""
        self.canel = False
        self.is_change = False
        self.is_only_verse = False

        self.document = Document()
        self.document.styles['Normal'].font.name = 'Times New Roman'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')







        self.top = tkinter.Tk()
        self.top.protocol('WM_DELETE_WINDOW', self.closeWindow)
        self.data_item = {}
        self.top.title("经文范围生成工具")
        self.top.geometry('600x480+600+200')


        self.docTk = tkinter.Tk()
        self.docTk.title("保存文件")
        self.docTk.geometry('300x200+750+340')
        self.docTk.minsize(300, 200)  # 最小尺寸

        self.docTk.maxsize(300, 200)  # 最大尺寸
        self.docTk.resizable(width=False, height=False)

        self.zxjd = Label(self.docTk, font=("宋体", 15))
        self.zxjd.place(x=240, y=106)
        self.wdscjd = Label(self.docTk, text="文档生成进度", font=("楷体", 18))
        self.wdscjd.place(x=70, y=30)
        self.file_day_var = StringVar()
        self.saveFileDay = Label(self.docTk, text="周一：", font=("宋体", 15))
        self.saveFileDay.place(x=20, y=108)


        self.canvas = Canvas(self.docTk, width=160, height=40, bg="white")
        self.canvas.place(x=75, y=100)
        self.out_rec = self.canvas.create_rectangle(5, 5, 155, 35, outline="blue", width=1)
        self.fill_rec = self.canvas.create_rectangle(5, 5, 5, 35, outline="", width=0, fill="blue")

        self.file_label = Label(self.docTk, text="请输入保存文件名称：", font=("楷体", 15))
        self.file_label.place(x=50, y=30)
        self.docname = Text(self.docTk, width=24, height=1, font=("楷体", 15), undo=True, wrap=NONE)
        self.docname.place(x=30, y=80)
        self.save_button = Button(self.docTk, text="点击确认", command=self.saveFile, font=("楷体", 15))
        self.save_button.place(x=100, y=120)
        self.save_button.place_forget()
        self.docname.place_forget()
        self.file_label.place_forget()

        self.top.minsize(640, 480)  # 最小尺寸

        self.top.maxsize(640, 480)  # 最大尺寸
        self.top.resizable(width=False, height=False)
        self.this_day_var = StringVar()
        self.this_day = Label(self.top,font=("楷体",18),textvariable=self.this_day_var).place(x=280, y=80)
        Label(self.top, text="课程名称：", font=("楷体", 15)).place(x=70, y=30)
        self.kcmc_var = StringVar()
        test_cmd = self.top.register(self.findCourse)
        self.kcmc = Entry(self.top, width=35, font=("楷体", 15), textvariable=self.kcmc_var,validate="key",validatecommand=(test_cmd, '%P', '%v', '%W'))
        self.kcmc.place(x=170, y=30)
        self.this_day_var.set("周一")
        Label(self.top, text="课程内容：", font=("楷体", 15)).place(x=70, y=130)
        self.kcnr_var = StringVar()
        self.kcnr = Entry(self.top, width=35, font=("楷体", 15), textvariable=self.kcnr_var).place(x=170, y=130)
        Label(self.top, text="叙述经文：", font=("楷体", 15)).place(x=70, y=180)
        self.xsjw = Text(self.top, width=35, height=5,font=("楷体",15))
        self.xsjw.place(x=170, y=180)
        self.xsjw.bind("<Alt-q>",self.KeyPress)
        Label(self.top, text="读经进度：", font=("楷体", 15)).place(x=70, y=330)
        self.djjd_var = StringVar()
        self.djjd = Entry(self.top, width=35, font=("楷体", 15),textvariable=self.djjd_var)
        self.djjd.place(x=170, y=330)
        self.djjd_var.set("无")
        self.djjd.bind("<Alt-w>", self.KeyPress1)
        self.button1 = Button(self.top, text="上一页", command=self.prev, font=("楷体", 15))
        self.button1.place(x=200, y=400)
        self.button2 = Button(self.top, text="下一页", command=self.next, font=("楷体", 15))
        self.button2.place(x=360, y=400)
        self.button3 = Button(self.top, text="确定", command=self.getDocx, font=("楷体", 15))
        self.button3.place(x=360, y=400)
        self.button3.place_forget()
        self.button1.place_forget()

        # self.top.bind('<KeyPress-Up>',self.prev)
        self.top.bind('<KeyPress-Down>', self.KeyPressDown)
        # self.top.bind('<Return>', self.getDocx)
        self.docTk.withdraw()

        self.top.mainloop()
        self.docTk.mainloop()

    def findCourse(self,content, reason, name):
        self.search_list = self.w_search(content)
        json = {"周一":0,"周二":1,"周三":2,"周四":3,"周五":4,"周六":5,"主日":6}
        if(len(self.search_list)!=0):
            self.kcnr_var.set(self.search_list[json[self.this_day_var.get()]])
        return True

    def KeyPressUp(self, event):
        self.prev()

    def KeyPressDown(self, event):
        self.next()

    def KeyPressReturn(self, event):
        self.getDocx()


    def KeyPress(self,event):
        text = self.xsjw.get('0.0',END)
        split_number = {
            '创世记': '创世记', '出埃及记': '出埃及记', '利未记': '利未记', '民数记': '民数记', '申命记': '申命记', '约书亚记': '约书亚记', '士师记': '士师记',
             '路得记': '路得记', '撒母耳记上': '撒母耳记上', '撒母耳记下': '撒母耳记下', '列王记上': '列王记上', '列王记下': '列王记下', '历代志上': '历代志上',
             '历代志下': '历代志下', '以斯拉记': '以斯拉记', '尼希米记': '尼希米记', '以斯帖记': '以斯帖记', '约伯记': '约伯记', '诗篇': '诗篇', '箴言': '箴言',
             '传道书': '传道书', '雅歌': '雅歌', '以赛亚书': '以赛亚书', '耶利米书': '耶利米书', '耶利米哀歌': '耶利米哀歌', '以西结书': '以西结书', '但以理书': '但以理书',
             '何西阿书': '何西阿书', '约珥书': '约珥书', '阿摩司书': '阿摩司书', '俄巴底亚书': '俄巴底亚书', '约拿书': '约拿书', '弥迦书': '弥迦书', '那鸿书': '那鸿书',
             '哈巴谷书': '哈巴谷书', '西番雅书': '西番雅书', '哈该书': '哈该书', '撒迦利亚书': '撒迦利亚书', '玛拉基书': '玛拉基书', '马太福音': '马太福音',
             '马可福音': '马可福音', '路加福音': '路加福音', '约翰福音': '约翰福音', '使徒行传': '使徒行传', '罗马书': '罗马书', '哥林多前书': '哥林多前书',
             '哥林多后书': '哥林多后书', '加拉太书': '加拉太书', '以弗所书': '以弗所书', '腓立比书': '腓立比书', '歌罗西书': '歌罗西书', '帖撒尼罗迦前书': '帖撒尼罗迦前书',
             '帖撒尼罗迦后书': '帖撒尼罗迦后书', '提摩太前书': '提摩太前书', '提摩太后书': '提摩太后书', '提多书': '提多书', '腓利门书': '腓利门书', '希伯来书': '希伯来书',
             '雅各书': '雅各书', '彼得前书': '彼得前书', '彼得后书': '彼得后书', '约翰一书': '约翰一书', '约翰二书': '约翰二书', '约翰三书': '约翰三书', '犹大书': '犹大书',
             '启示录': '启示录',

            '出埃及': '出埃及记', '利未': '利未记', '民数': '民数记', '申命': '申命记', '约书亚': '约书亚记', '士师': '士师记',
            '路得': '路得记', '撒母耳上': '撒母耳记上', '撒母耳下': '撒母耳记下', '列王上': '列王记上', '列王下': '列王记下', '历代上': '历代志上',
            '历代下': '历代志下', '以斯拉': '以斯拉记', '尼希米': '尼希米记', '以斯帖': '以斯帖记', '约伯': '约伯记',
            '传道': '传道书', '以赛亚': '以赛亚书', '耶利米': '耶利米书', '以西结': '以西结书', '但以理': '但以理书',
            '何西阿': '何西阿书', '约珥': '约珥书', '阿摩司': '阿摩司书', '俄巴底亚': '俄巴底亚书', '约拿': '约拿书', '弥迦': '弥迦书', '那鸿': '那鸿书',
            '哈巴谷': '哈巴谷书', '西番雅': '西番雅书', '哈该': '哈该书', '撒迦利亚': '撒迦利亚书', '玛拉基': '玛拉基书',
            '约翰': '约翰福音', '使徒': '使徒行传', '罗马': '罗马书', '哥林多前': '哥林多前书',
            '哥林多后': '哥林多后书', '加拉太': '加拉太书', '以弗所': '以弗所书', '腓立比': '腓立比书', '歌罗西': '歌罗西书', '帖撒尼罗迦前': '帖撒尼罗迦前书',
            '帖撒尼罗迦后': '帖撒尼罗迦后书', '提摩太前': '提摩太前书', '提摩太后': '提摩太后书', '提多': '提多书', '腓利门': '腓利门书', '希伯来': '希伯来书',
            '雅各': '雅各书', '彼得前': '彼得前书', '彼得后': '彼得后书', '犹大': '犹大书',


        '约一': '约翰一书', '约二': '约翰二书', '约三': '约翰三书','约壹': '约翰一书', '约贰': '约翰二书', '约叁': '约翰三书',
                          '撒上': '撒母耳记上', '撒下': '撒母耳记下', '王上': '列王记上', '王下': '列王记下', '代上': '历代志上',
                          '代下': '历代志下', '提后': '提摩太后书', '马可': '马可福音','路加': '路加福音', '马太': '马太福音', '可': '马可福音',
                         '创': '创世记', '出': '出埃及记', '利': '利未记', '民': '民数记', '申': '申命记', '书': '约书亚记', '士': '士师记',
                         '得': '路得记',  '拉': '以斯拉记', '尼': '尼希米记', '斯': '以斯帖记', '伯': '约伯记', '诗': '诗篇', '箴': '箴言',
                         '传': '传道书', '歌': '雅歌', '赛': '以赛亚书', '耶': '耶利米书', '哀': '耶利米哀歌', '结': '以西结书', '但': '但以理书',
                         '何': '何西阿书', '珥': '约珥书', '摩': '阿摩司书', '俄': '俄巴底亚书', '拿': '约拿书', '弥': '弥迦书', '鸿': '那鸿书',
                         '哈': '哈巴谷书', '番': '西番雅书', '该': '哈该书', '亚': '撒迦利亚书', '玛': '玛拉基书', '太': '马太福音',
                         '路': '路加福音', '约': '约翰福音', '徒': '使徒行传', '罗': '罗马书', '林前': '哥林多前书', '林后': '哥林多后书', '加': '加拉太书',
                         '弗': '以弗所书', '腓': '腓立比书', '西': '歌罗西书', '帖前': '帖撒尼罗迦前书', '帖后': '帖撒尼罗迦后书', '提前': '提摩太前书',
                         '多': '提多书', '门': '腓利门书', '来': '希伯来书', '雅': '雅各书', '彼前': '彼得前书', '彼后': '彼得后书',
                         '犹': '犹大书', '启': '启示录',":":"章","至":"~","，":"节，",",":"节，",
                            "一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7", "八": "8", "九": "9", "壹": "1","-":"~",
                            "十":"10","。":"节，"
                        }
        dict_num = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7", "八": "8", "九": "9", "壹": "1",
                     "十":"10",     }

        text1 = ""
        while(True):
            isHave = False
            xlen = 1
            for k in split_number:
                split_str = text[:len(k)]
                if split_str in split_number:
                    xlen = len(split_str)
                    isHave = True
                    break


            if(isHave is True):
                if (text[:xlen] in (",","，")) and text1[-1] != '节' :
                    text1 += split_number[text[:xlen]]
                elif (text[:xlen] in (",","，")):
                    text1 += text[:xlen]
                elif text[:xlen] == "。" and xlen == len(text) - 1 and text1[-1] != '节':
                    text1 += "节" + text[:xlen]
                elif(text[:xlen]=="。" and xlen!=len(text)-1):
                    text1 += split_number[text[:xlen]]
                elif text[:xlen]=="。" and xlen==len(text)-1:
                    text1 += text[:xlen]
                elif text[:xlen] in dict_num and text[:xlen]!="十":
                    if text[xlen:xlen+1]!='章' and (text[xlen:xlen+1] in split_number and split_number[text[xlen:xlen+1]].isdigit() ):
                        text1 += split_number[text[:xlen]]
                    else:
                        text1 += split_number[text[:xlen]]+"章"
                elif(text[:xlen]=="十"):
                    if(text1[-1].isdigit() and (text[xlen:xlen+1] in split_number) and split_number[text[xlen:xlen+1]].isdigit()):
                        text1 += ""
                    elif text1[-1].isdigit() and text[xlen:xlen+1]!='章':
                        text1 += "0章"
                    elif text1[-1].isdigit() and text[xlen:xlen+1]=='章':
                        text1 += "0"
                    elif (text[xlen:xlen+1] in split_number) and split_number[text[xlen:xlen+1]].isdigit():
                        text1 += "1"
                    elif text[xlen:xlen+1]!='章':
                        text1 += split_number[text[:xlen]]+"章"
                    else:
                        text1 += split_number[text[:xlen]]
                else:
                    text1 += split_number[text[:xlen]]
                text = text[xlen:]

            else:
                text1 += text[:1]
                text = text[1:]
            if(len(text) <1):
                break

        self.xsjw.delete(0.0, END)
        self.xsjw.insert(END, text1)


    def KeyPress1(self,event):
        text =self.djjd_var.get()
        split_number = {
            '创世记': '创世记', '出埃及记': '出埃及记', '利未记': '利未记', '民数记': '民数记', '申命记': '申命记', '约书亚记': '约书亚记', '士师记': '士师记',
             '路得记': '路得记', '撒母耳记上': '撒母耳记上', '撒母耳记下': '撒母耳记下', '列王记上': '列王记上', '列王记下': '列王记下', '历代志上': '历代志上',
             '历代志下': '历代志下', '以斯拉记': '以斯拉记', '尼希米记': '尼希米记', '以斯帖记': '以斯帖记', '约伯记': '约伯记', '诗篇': '诗篇', '箴言': '箴言',
             '传道书': '传道书', '雅歌': '雅歌', '以赛亚书': '以赛亚书', '耶利米书': '耶利米书', '耶利米哀歌': '耶利米哀歌', '以西结书': '以西结书', '但以理书': '但以理书',
             '何西阿书': '何西阿书', '约珥书': '约珥书', '阿摩司书': '阿摩司书', '俄巴底亚书': '俄巴底亚书', '约拿书': '约拿书', '弥迦书': '弥迦书', '那鸿书': '那鸿书',
             '哈巴谷书': '哈巴谷书', '西番雅书': '西番雅书', '哈该书': '哈该书', '撒迦利亚书': '撒迦利亚书', '玛拉基书': '玛拉基书', '马太福音': '马太福音',
             '马可福音': '马可福音', '路加福音': '路加福音', '约翰福音': '约翰福音', '使徒行传': '使徒行传', '罗马书': '罗马书', '哥林多前书': '哥林多前书',
             '哥林多后书': '哥林多后书', '加拉太书': '加拉太书', '以弗所书': '以弗所书', '腓立比书': '腓立比书', '歌罗西书': '歌罗西书', '帖撒尼罗迦前书': '帖撒尼罗迦前书',
             '帖撒尼罗迦后书': '帖撒尼罗迦后书', '提摩太前书': '提摩太前书', '提摩太后书': '提摩太后书', '提多书': '提多书', '腓利门书': '腓利门书', '希伯来书': '希伯来书',
             '雅各书': '雅各书', '彼得前书': '彼得前书', '彼得后书': '彼得后书', '约翰一书': '约翰一书', '约翰二书': '约翰二书', '约翰三书': '约翰三书', '犹大书': '犹大书',
             '启示录': '启示录',

            '出埃及': '出埃及记', '利未': '利未记', '民数': '民数记', '申命': '申命记', '约书亚': '约书亚记', '士师': '士师记',
            '路得': '路得记', '撒母耳上': '撒母耳记上', '撒母耳下': '撒母耳记下', '列王上': '列王记上', '列王下': '列王记下', '历代上': '历代志上',
            '历代下': '历代志下', '以斯拉': '以斯拉记', '尼希米': '尼希米记', '以斯帖': '以斯帖记', '约伯': '约伯记',
            '传道': '传道书', '以赛亚': '以赛亚书', '耶利米': '耶利米书', '以西结': '以西结书', '但以理': '但以理书',
            '何西阿': '何西阿书', '约珥': '约珥书', '阿摩司': '阿摩司书', '俄巴底亚': '俄巴底亚书', '约拿': '约拿书', '弥迦': '弥迦书', '那鸿': '那鸿书',
            '哈巴谷': '哈巴谷书', '西番雅': '西番雅书', '哈该': '哈该书', '撒迦利亚': '撒迦利亚书', '玛拉基': '玛拉基书',
            '约翰': '约翰福音', '使徒': '使徒行传', '罗马': '罗马书', '哥林多前': '哥林多前书',
            '哥林多后': '哥林多后书', '加拉太': '加拉太书', '以弗所': '以弗所书', '腓立比': '腓立比书', '歌罗西': '歌罗西书', '帖撒尼罗迦前': '帖撒尼罗迦前书',
            '帖撒尼罗迦后': '帖撒尼罗迦后书', '提摩太前': '提摩太前书', '提摩太后': '提摩太后书', '提多': '提多书', '腓利门': '腓利门书', '希伯来': '希伯来书',
            '雅各': '雅各书', '彼得前': '彼得前书', '彼得后': '彼得后书', '犹大': '犹大书',


        '约一': '约翰一书', '约二': '约翰二书', '约三': '约翰三书','约壹': '约翰一书', '约贰': '约翰二书', '约叁': '约翰三书',
                          '撒上': '撒母耳记上', '撒下': '撒母耳记下', '王上': '列王记上', '王下': '列王记下', '代上': '历代志上',
                          '代下': '历代志下', '提后': '提摩太后书', '马可': '马可福音','路加': '路加福音', '马太': '马太福音', '可': '马可福音',
                         '创': '创世记', '出': '出埃及记', '利': '利未记', '民': '民数记', '申': '申命记', '书': '约书亚记', '士': '士师记',
                         '得': '路得记',  '拉': '以斯拉记', '尼': '尼希米记', '斯': '以斯帖记', '伯': '约伯记', '诗': '诗篇', '箴': '箴言',
                         '传': '传道书', '歌': '雅歌', '赛': '以赛亚书', '耶': '耶利米书', '哀': '耶利米哀歌', '结': '以西结书', '但': '但以理书',
                         '何': '何西阿书', '珥': '约珥书', '摩': '阿摩司书', '俄': '俄巴底亚书', '拿': '约拿书', '弥': '弥迦书', '鸿': '那鸿书',
                         '哈': '哈巴谷书', '番': '西番雅书', '该': '哈该书', '亚': '撒迦利亚书', '玛': '玛拉基书', '太': '马太福音',
                         '路': '路加福音', '约': '约翰福音', '徒': '使徒行传', '罗': '罗马书', '林前': '哥林多前书', '林后': '哥林多后书', '加': '加拉太书',
                         '弗': '以弗所书', '腓': '腓立比书', '西': '歌罗西书', '帖前': '帖撒尼罗迦前书', '帖后': '帖撒尼罗迦后书', '提前': '提摩太前书',
                         '多': '提多书', '门': '腓利门书', '来': '希伯来书', '雅': '雅各书', '彼前': '彼得前书', '彼后': '彼得后书',
                         '犹': '犹大书', '启': '启示录',":":"章","至":"~","，":"节，",",":"节，",
                            "一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7", "八": "8", "九": "9", "壹": "1","-":"~",
                            "十":"10","。":"节，"
                        }
        dict_num = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7", "八": "8", "九": "9", "壹": "1",
                     "十":"10",     }

        text1 = ""
        while(True):
            isHave = False
            xlen = 1
            for k in split_number:
                split_str = text[:len(k)]
                if split_str in split_number:
                    xlen = len(split_str)
                    isHave = True
                    break


            if(isHave is True):
                if (text[:xlen] in (",","，")) and text1[-1] != '节' :
                    text1 += split_number[text[:xlen]]
                elif (text[:xlen] in (",","，")):
                    text1 += text[:xlen]
                elif text[:xlen] == "。" and xlen == len(text) - 1 and text1[-1] != '节':
                    text1 += "节" + text[:xlen]
                elif(text[:xlen]=="。" and xlen!=len(text)-1):
                    text1 += split_number[text[:xlen]]
                elif text[:xlen]=="。" and xlen==len(text)-1:
                    text1 += text[:xlen]
                elif text[:xlen] in dict_num and text[:xlen]!="十":
                    if text[xlen:xlen+1]!='章' and (text[xlen:xlen+1] in split_number and split_number[text[xlen:xlen+1]].isdigit() ):
                        text1 += split_number[text[:xlen]]
                    else:
                        text1 += split_number[text[:xlen]]+"章"
                elif(text[:xlen]=="十"):
                    if(text1[-1].isdigit() and (text[xlen:xlen+1] in split_number) and split_number[text[xlen:xlen+1]].isdigit()):
                        text1 += ""
                    elif text1[-1].isdigit() and text[xlen:xlen+1]!='章':
                        text1 += "0章"
                    elif text1[-1].isdigit() and text[xlen:xlen+1]=='章':
                        text1 += "0"
                    elif (text[xlen:xlen+1] in split_number) and split_number[text[xlen:xlen+1]].isdigit():
                        text1 += "1"
                    elif text[xlen:xlen+1]!='章':
                        text1 += split_number[text[:xlen]]+"章"
                    else:
                        text1 += split_number[text[:xlen]]
                else:
                    text1 += split_number[text[:xlen]]
                text = text[xlen:]

            else:
                text1 += text[:1]
                text = text[1:]
            if(len(text) <1):
                break

        self.djjd_var.set(text1)

    def closeWindow(self):
        ans = tkinter.messagebox.askyesno(title='关闭程序',message='确定要关闭程序吗?')
        if(ans):
            sys.exit()
        else:
            return

    # 验证是否有效
    def valid(self,bible_volume_name, bible_verse_node_str, crt_volume, crt_verse):
        is_valid = True
        # 卷名+章+节格式
        if bible_volume_name != "":
            crt_volume = bible_volume_name
            # 输入卷名不在bible_dict中代表输入错误
            if bible_volume_name not in self.bible_dict.keys():
                is_valid = False
                return is_valid, crt_volume, crt_verse
            split_flag = ""
            if bible_volume_name == "诗篇":
                split_flag = "篇"
            else:
                split_flag = "章"
            # 分割章节
            verse = bible_verse_node_str.split(split_flag)

            if len(verse) == 2:
                crt_verse = verse[0]
            if (len(verse)>2):
                is_valid = False
                return is_valid, crt_volume, crt_verse
            # 分割后章节都存在的情况
            if len(verse) == 2 and verse[1] != "" and verse[1][-1] != "" and verse[1][-1] != "。":
                if(len(verse[1].split("节"))==2 and verse[1].split("节")[1]!=""):
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if verse[0].isdigit() is False or verse[1].split("节")[0].isdigit() is False:
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if int(verse[0]) > len(self.valid_dict[crt_volume]) or int(verse[0]) <= 0:
                    tkinter.messagebox.showerror("章数错误","检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
                crt_verse = verse[0]
                if (int(verse[1].split("节")[0]) > self.valid_dict[crt_volume][self.dealVerseName(verse[0])] or int(
                    verse[1].split("节")[0]) <= 0):
                    tkinter.messagebox.showerror("节数错误","检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
        # 章+节  章   节 等情况
        else:
            split_flag = ""
            if crt_volume == "":
                is_valid = False
                return is_valid, crt_volume
            if crt_volume == "诗篇":
                split_flag = "篇"
            else:
                split_flag = "章"
            verse = bible_verse_node_str.split(split_flag)
            if len(verse) == 2:
                crt_verse = verse[0]
            # 章+节情况
            if len(verse) == 2 and verse[1] != "" and verse[1][-1] != "" and verse[1][-1] != "。":
                if (len(verse[1].split("节"))==2 and verse[1].split("节")[1] != ""):
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if verse[0].isdigit() is False or verse[1].split("节")[0].isdigit() is False:
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if verse[0] != "" and int(verse[0]) > len(self.valid_dict[crt_volume]) or int(verse[0]) <= 0:
                    tkinter.messagebox.showerror("章数错误","检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
                crt_verse = verse[0]
                if (int(verse[1].split("节")[0]) > self.valid_dict[crt_volume][self.dealVerseName(verse[0])] or int(
                    verse[1].split("节")[0]) <= 0):
                    tkinter.messagebox.showerror("节数错误","检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
            # 章
            elif len(verse) == 2 and (verse[1] == "" or verse[1] == "。"):
                if verse[0].isdigit() is False:
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if int(verse[0]) > len(self.valid_dict[crt_volume]) or int(verse[0]) <= 0:
                    tkinter.messagebox.showerror("章数错误","检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
                crt_verse = verse[0]
            # 节
            else:
                if (len(verse[0].split("节"))==2 and verse[0].split("节")[1] != ""):
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if verse[0].split("节")[0].isdigit() is False:
                    is_valid = False
                    return is_valid, crt_volume, crt_verse
                if (int(verse[0].split("节")[0]) > self.valid_dict[crt_volume][self.dealVerseName(crt_verse)] or int(
                    verse[0].split("节")[0]) <= 0):
                    tkinter.messagebox.showerror("节数错误","检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！")
                    is_valid = "2"
                    return is_valid, crt_volume, crt_verse
        return is_valid, crt_volume, crt_verse


    def bible_valid(self,bible_scope):
        crt_volume = ""
        crt_verse = ""
        is_valid = True
        if "。" in bible_scope:
            bible_scope = bible_scope[0:-1]
        scopes = re.split("[,]|[，]", bible_scope)
        if bible_scope != "无":
            for scope in scopes:
                before = ""
                after = ""
                if "~" in scope and len(scope.split("~")) == 2:
                    before, after = scope.split("~")
                elif "～" in scope and len(scope.split("～")) == 2:
                    before, after = scope.split("～")
                else:
                    before = scope
                    after = scope
                try:
                    bible_is_valid, new_crt_volume, new_crt_verse = self.valid(self.dealScope(before)[0], self.dealScope(before)[1],
                                                                          crt_volume, crt_verse)
                    crt_volume = new_crt_volume
                    crt_verse = new_crt_verse
                    if bible_is_valid is False:
                        is_valid = False
                        break
                    if bible_is_valid == "2":
                        is_valid = "1"
                        break

                    bible_is_valid, new_crt_volume1, new_crt_verse1 = self.valid(self.dealScope(after)[0], self.dealScope(after)[1],
                                                                            crt_volume, crt_verse)
                    crt_volume = new_crt_volume1
                    crt_verse = new_crt_verse1
                    if bible_is_valid is False:
                        is_valid = False
                        break

                    if bible_is_valid == "2":
                        is_valid = "1"
                        break
                except:
                    is_valid = False

        else:
            is_valid = "0"
        return is_valid


    # 删除文档段落
    def delete_paragraph(self,paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        # p._p = p._element = None
        paragraph._p = paragraph._element = None


    # 将输入分割为卷名和章节
    def dealScope(self,scope):
        bible_volume_name = ""
        bible_verse_node_str = ""
        is_number = False
        for idx in range(0, len(scope)):
            for n in range(0, len(self.number_list)):
                if scope[idx] == self.number_list[n]:
                    is_number = True
                    break
            if is_number is False:
                bible_volume_name += scope[idx]
            else:
                bible_verse_node_str += scope[idx]
        return bible_volume_name, bible_verse_node_str


    # 判断是否为节的情况
    def node_is_none(self,scope):
        is_node = False
        if len(scope) == 1:
            if scope[-1] == "节":
                is_node = True
        else:
            if scope[-1] == "节" or scope[-2] == "节":
                is_node = True
        for n in range(0, len(self.number_list)):
            if scope[-1] == self.number_list[n]:
                is_node = True
                break

        return is_node


    # 处理章数
    def dealVerseName(self,verse_name):
        verse_handle_name = ""
        if len(verse_name) == 1:
            verse_handle_name = self.num_dict[verse_name]
        elif len(verse_name) == 2:
            if verse_name[0] == "1":
                verse_handle_name = "十" + self.num_dict[verse_name[1]]
            else:
                verse_handle_name = self.num_dict[verse_name[0]] + "十" + self.num_dict[verse_name[1]]
        else:
            if verse_name[1] == "0" and verse_name[2] == "0":
                verse_handle_name = "一○○"
            elif verse_name[1] == "0":
                verse_handle_name = self.num_dict[verse_name[0]] + "○" + self.num_dict[verse_name[2]]
            elif verse_name[2] == "0":
                verse_handle_name = self.num_dict[verse_name[0]] + self.num_dict[verse_name[1]] + "○"
            else:
                verse_handle_name = self.num_dict[verse_name[0]] + self.num_dict[verse_name[1]] + self.num_dict[verse_name[2]]
        return verse_handle_name


    # dict根据val反取key
    def get_key(self,d, value):
        return [k for k, v in d.items() if v == value]


    # 处理写入的文章段落
    def p_format(self,p):
        paragraph_format = p.paragraph_format
        paragraph_format.space_before = Pt(0)  # 上行间距
        paragraph_format.space_after = Pt(0)  # 下行间距
        paragraph_format.line_spacing = Pt(0)  # 行距


    # 根据输入课程内容检索每周课程
    def w_search(self,c_input):
        input_split = re.split("[ ]|[　]", c_input)
        if(len(input_split)==2):
            c_input = input_split[0] + "　" + input_split[1]
        elif(len(input_split)==3):
            c_input = input_split[0] + "　" + input_split[2]
            # 第六课　大卫受试验以被成全
        else:
            return []
        search_list = []
        index = 0
        for paragraph in self.bible_travel_data.paragraphs:
            index = index + 1
            if paragraph.text.strip() == c_input:

                for idx in range(0, 7):
                    search_list.append(self.bible_travel_data.paragraphs[index + idx].text.strip())
        return search_list


    def vaild_bible(self):
        kcmc_val = self.kcmc_var.get()
        if kcmc_val == "":
            tkinter.messagebox.showerror("课程名称错误","课程名称不能为空，请重新输入！")
            return False
        if len(re.split("[ ]|[　]", kcmc_val)) == 1:
            tkinter.messagebox.showerror("课程名称错误","检测到课程名称输入有误，请重新输入！")
            return False

        kcnr_val = self.kcnr_var.get()
        if kcnr_val == "":
            tkinter.messagebox.showerror("课程内容错误","课程内容不能为空，请重新输入！")
            return False

        scope = self.xsjw.get(0.0, END).strip()
        if scope == "":
            tkinter.messagebox.showerror("叙述经文错误", "叙述经文不能为空！")
            return False
        scope_is_valid = self.bible_valid(scope)
        if scope_is_valid is False:
            tkinter.messagebox.showerror("叙述经文错误", "检测到叙述经文输入有误，请重新输入！")
            return False
        elif scope_is_valid=="1":
            return False
        return True

    def handle_bible(self):
        # try:
        self.docTk.deiconify()
        kcmc_val = self.kcmc_var.get()
        # 课程名称
        row_one = self.document.add_paragraph(kcmc_val)


        for item in self.data_item:
            # 进度名称变换
            self.saveFileDay.config(text=item)
            #row_two 课程内容名称
            row_two = self.document.add_paragraph(self.data_item[item]["课程内容"])
            scope = self.data_item[item]['叙述经文'].strip()
            row_three = self.document.add_paragraph()
            title = row_three.add_run("叙述经文：")
            title.font.color.rgb = RGBColor(255, 0, 0)
            title.bold = True
            # row_three 叙述经文 经文范围
            content = row_three.add_run(scope)
            # 颜色
            content.font.color.rgb = RGBColor(255, 0, 0)
            # 高亮颜色
            content.font.highlight_color = MSO_THEME_COLOR_INDEX.ACCENT_3
            # 加粗
            content.bold = True
            # 加入到文档中
            self.p_format(row_one)
            self.p_format(row_two)
            self.p_format(row_three)
            # 经文范围先去掉末尾的句号，方便做处理
            if "。" in scope:
                scope = scope[0:-1]
            # 经文范围按照逗号分隔
            bible_read_list = re.split("[,]|[，]", scope)
            # 进度条归0
            self.progress_bar1(0 / len(bible_read_list) ,1)
            # 每一段的经文范围都要做处理
            for i in range(0, len(bible_read_list)):
                # 进度条更新
                self.progress_bar1((i + 1 / 10) / len(bible_read_list), 1)
                # 前卷名
                bible_before_volume_name = ""
                # 前章名
                bible_before_verse_name = ""
                # 前节名
                bible_before_node_name = ""
                # 后卷名
                bible_after_volume_name = ""
                # 后章名
                bible_after_verse_name = ""
                # 后节名
                bible_after_node_name = ""
                # 经文范围按照~分隔
                if "~" in bible_read_list[i]:
                    scope_before = bible_read_list[i].split("~")
                else:
                    scope_before = bible_read_list[i].split("～")
                # 有范围  后与前各取，无范围，后与前相同
                if len(scope_before) == 2:
                    bible_volume_before, bible_verse_node_before = self.dealScope(scope_before[0])
                    bible_volume_after, bible_verse_node_after = self.dealScope(scope_before[1])
                else:
                    bible_volume_before, bible_verse_node_before = self.dealScope(scope_before[0])
                    bible_volume_after, bible_verse_node_after = self.dealScope(scope_before[0])
                if len(scope_before) == 2:
                    # 有章节且卷名不为空的情况
                    if bible_volume_before != "":
                        bible_before_volume_name = bible_volume_before
                        self.current_volume_name = bible_before_volume_name
                    else:
                        bible_before_volume_name = self.current_volume_name
                    split_flag = ""
                    if bible_before_volume_name != "诗篇":
                        split_flag = "章"
                    else:
                        split_flag = "篇"
                    # 卷名+节情况
                    if (len(bible_verse_node_before.split(split_flag)) == 2 and
                        bible_verse_node_before.split(split_flag)[1] != "" and
                        bible_verse_node_before.split(split_flag)[1] != "。"):
                        bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                        self.current_verse = bible_before_verse_name
                        bible_before_node_name = bible_verse_node_before.split(split_flag)[1].split("节")[0]
                    else:
                        is_node = self.node_is_none(bible_verse_node_before)
                        after_is_node = self.node_is_none(bible_verse_node_after)
                        # 只有节
                        if is_node and after_is_node:
                            bible_before_verse_name = self.current_verse
                            bible_before_node_name = bible_verse_node_before.split("节")[0]
                        # 只有章
                        else:
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            bible_before_node_name = "1"
                    # 与前对应
                    if bible_volume_after != "":
                        bible_after_volume_name = bible_volume_after
                        # 章+节
                        if (len(bible_verse_node_after.split(split_flag)) == 2 and
                            bible_verse_node_after.split(split_flag)[1] != "" and
                            bible_verse_node_after.split(split_flag)[1] != "。"):
                            bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                            bible_after_node_name = bible_verse_node_after.split(split_flag)[1].split("节")[0]
                        else:
                            is_node = self.node_is_none(bible_verse_node_after)
                            # 只有节
                            if is_node:
                                bible_after_verse_name = self.current_verse
                                bible_after_node_name = bible_verse_node_after.split("节")[0]
                            # 只有章
                            else:
                                bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                                bible_after_node_name = "999"
                    # 无卷名情况
                    else:
                        bible_after_volume_name = self.current_volume_name
                        # 章+节
                        if (len(bible_verse_node_after.split(split_flag)) == 2 and
                            bible_verse_node_after.split(split_flag)[1] != "" and
                            bible_verse_node_after.split(split_flag)[1] != "。"):
                            bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                            self.current_verse = bible_after_verse_name
                            bible_after_node_name = bible_verse_node_after.split(split_flag)[1].split("节")[0]
                        else:
                            is_node = self.node_is_none(bible_verse_node_after)
                            # 节
                            if is_node:
                                bible_after_verse_name = self.current_verse
                                bible_after_node_name = bible_verse_node_after.split("节")[0]
                            # 章
                            else:
                                bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                                bible_after_node_name = "999"
                # 后与前相同
                else:
                    # 卷名+章节
                    if bible_volume_before != "":
                        if bible_volume_before != "诗篇":
                            split_flag = "章"
                        else:
                            split_flag = "篇"
                        if (len(bible_verse_node_before.split(split_flag)) == 2 and
                            bible_verse_node_before.split(split_flag)[1] != "" and
                            bible_verse_node_before.split(split_flag)[1] != "。"):
                            bible_before_volume_name = bible_volume_before
                            self.current_volume_name = bible_before_volume_name
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            self.current_verse = bible_before_verse_name
                            bible_before_node_name = bible_verse_node_before.split(split_flag)[1].split("节")[0]
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = bible_before_node_name
                        else:
                            bible_before_volume_name = bible_volume_before
                            self.current_volume_name = bible_before_volume_name
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            self.current_verse = bible_before_verse_name
                            bible_before_node_name = "1"
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = "999"
                    # 章节
                    else:
                        bible_before_volume_name = self.current_volume_name
                        if bible_before_volume_name != "诗篇":
                            split_flag = "章"
                        else:
                            split_flag = "篇"
                        if (len(bible_verse_node_before.split(split_flag)) == 2 and
                            bible_verse_node_before.split(split_flag)[1] != "" and
                            bible_verse_node_before.split(split_flag)[1] != "。"):
                            bible_before_verse_name = scope_before[0].split(split_flag)[0]
                            self.current_verse = bible_before_verse_name
                            bible_before_node_name = scope_before[0].split(split_flag)[1].split("节")[0]
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = bible_before_node_name
                        else:
                            is_node = self.node_is_none(bible_verse_node_before)
                            # 节
                            if is_node:
                                bible_before_verse_name = self.current_verse
                                bible_before_node_name = bible_verse_node_before.split("节")[0]
                                bible_after_volume_name = bible_before_volume_name
                                bible_after_verse_name = bible_before_verse_name
                                bible_after_node_name = bible_before_node_name
                            # 章
                            else:
                                bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                                self.current_verse = bible_before_verse_name
                                bible_before_node_name = "1"
                                bible_after_volume_name = bible_before_volume_name
                                bible_after_verse_name = bible_before_verse_name
                                bible_after_node_name = "999"
                before_verse_handle_name = ""
                after_verse_handle_name = ""
                before_verse_handle_name = self.dealVerseName(bible_before_verse_name)
                after_verse_handle_name = self.dealVerseName(str(int(bible_after_verse_name) + 1))
                if bible_before_volume_name != "诗篇":
                    bible_before_volume_verse = bible_before_volume_name + "第" + before_verse_handle_name + "章"
                else:
                    bible_before_volume_verse = bible_before_volume_name + "第" + before_verse_handle_name + "篇"
                if bible_after_volume_name != "诗篇":
                    bible_after_volume_verse = bible_after_volume_name + "第" + after_verse_handle_name + "章"
                else:
                    bible_after_volume_verse = bible_after_volume_name + "第" + after_verse_handle_name + "篇"
                begin_row = 0
                done_row = 0
                is_volume_end = True
                is_bible_end = True
                # 先计算读经范围是哪一章到哪一章
                self.progress_bar1((i + 2 / 10) / len(bible_read_list), 1)
                before_bar_count = 0
                for paragraph in self.bible_data.paragraphs:
                    begin_row = begin_row + 1
                    if paragraph.text.strip() == bible_before_volume_verse:
                        break
                    if self.row_dict[bible_before_volume_name][before_verse_handle_name] > 5000:
                        before_bar_num = self.row_dict[bible_before_volume_name][before_verse_handle_name] // 5000
                        if begin_row % 5000 == 0:
                            before_bar_count = before_bar_count + 1
                            self.progress_bar1((i + 2 / 10 + (2.5 / 10) * (before_bar_count / before_bar_num)) / len(
                                bible_read_list), 1)
                self.progress_bar1((i + 4.5 / 10) / len(bible_read_list), 1)
                after_bar_count = 0
                for paragraph in self.bible_data.paragraphs:
                    done_row = done_row + 1
                    if paragraph.text.strip() == bible_after_volume_verse:
                        is_volume_end = False
                        is_bible_end = False
                        break
                    after_bar_num = ""
                    if after_verse_handle_name in self.row_dict[bible_after_volume_name] and \
                        self.row_dict[bible_after_volume_name][after_verse_handle_name] > 5000:
                        after_bar_num = self.row_dict[bible_after_volume_name][after_verse_handle_name] // 5000
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            self.progress_bar1((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list), 1)
                    elif after_verse_handle_name == "二十三" and bible_after_volume_name == "启示录":
                        after_bar_num = self.row_dict[bible_after_volume_name]["二十二"] // 5000
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            self.progress_bar1((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list), 1)
                    else:
                        key = self.get_key(self.index_dict, bible_after_volume_name)[0]
                        after_bar_num = self.row_dict[self.index_dict[str(int(key) + 1)]]["一"]
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            self.progress_bar1((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list),1)
                self.progress_bar1((i + 7 / 10) / len(bible_read_list), 1)
                # 读经范围到启示录的情况
                if is_volume_end:
                    done_row = 0
                    key = self.get_key(self.index_dict, bible_after_volume_name)[0]
                    if int(key) + 1 > 66:
                        is_bible_end = True
                    else:
                        bible_after_volume_verse = self.index_dict[str(int(key) + 1)] + "第一章"
                        for paragraph in self.bible_data.paragraphs:
                            done_row = done_row + 1
                            if paragraph.text.strip() == bible_after_volume_verse:
                                is_bible_end = False
                                break
                # 如果是启示录，直接预设最大行数
                if is_bible_end:
                    done_row = 49999
                begin_node = begin_row
                end_node_num = begin_row
                end_node = end_node_num
                row = 0
                # 范围进一步精确到哪一章哪一节
                for paragraph in self.bible_data.paragraphs:
                    row = row + 1
                    if begin_row < row < done_row:
                        begin = bible_before_verse_name + ":" + bible_before_node_name
                        begin_re = r'' + begin + '.*'
                        match = re.match(begin_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            begin_node = begin_node + 1
                            break
                        else:
                            begin_node = begin_node + 1
                self.progress_bar1((i + 8 / 10) / len(bible_read_list), 1)
                row = 0
                for paragraph in self.bible_data.paragraphs:
                    row = row + 1
                    if begin_row < row < done_row:
                        end = bible_after_verse_name + ":" + bible_after_node_name
                        end_re = r'' + end + '.*'
                        match = re.match(end_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            end_node = end_node + 1
                            break
                        else:
                            end_node = end_node + 1
                row = 0
                volume_name = bible_before_volume_name
                self.progress_bar1((i + 9 / 10) / len(bible_read_list), 1)
                # 检索中间的段落，并写入新文档中
                for paragraph in self.bible_data.paragraphs:
                    row = row + 1
                    if begin_node <= row <= end_node:
                        verse_name_re = ""
                        if volume_name != "诗篇":
                            verse_name_re = r'[\u4E00-\u9FA5]+第[\u4E00-\u9FA5]+章'
                        else:
                            verse_name_re = r'[\u4E00-\u9FA5]+第[\u4E00-\u9FA5]+篇'
                        vn = re.match(verse_name_re, paragraph.text.strip(), re.I)
                        if vn is not None:
                            volume_name = paragraph.text.split("第")[0]
                        vb_re = r'[0-9]{1,3}:[0-9]{1,3}.*'
                        match = re.match(vb_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            head = paragraph.text.split("　")[0]
                            content = paragraph.text.split("　")[1]
                            deal_content = re.sub('[0-9]', "", content)
                            label = self.bible_dict[volume_name] + head
                            row_result = "　" + deal_content
                            p = self.document.add_paragraph()
                            run = p.add_run(label)
                            run.font.color.rgb = RGBColor(0, 112, 192)
                            run.bold = True
                            p.add_run(row_result)
                            self.p_format(p)



            p = self.document.add_paragraph()
            p_title = p.add_run("读经进度：")
            p_title.font.color.rgb = RGBColor(255, 0, 0)
            p_title.bold = True
            run = p.add_run(self.data_item[item]["读经进度"])
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            self.progress_bar1(1, 1)
        self.wdscjd.place_forget()
        self.zxjd.place_forget()
        self.canvas.place_forget()
        self.saveFileDay.place_forget()
        self.file_label.place(x=50, y=30)
        self.docname.place(x=30, y=80)
        self.save_button.place(x=100, y=120)



    def saveFile(self):
        docname = self.docname.get(0.0,END).strip()
        self.document.save(docname+".docx")
        self.docTk.withdraw()
        tkinter.messagebox.showinfo("保存文件", "保存成功!")
        for item in self.data_item:
            self.data_item[item]['叙述经文'] = ""
            self.data_item[item]['读经进度'] = "无"
        self.xsjw.delete(0.0, END)
        self.djjd_var.set("无")
        self.wdscjd.place(x=70, y=30)
        self.saveFileDay.place(x=20, y=108)
        self.canvas.place(x=75, y=100)
        self.zxjd.place(x=240, y=106)
        self.file_label.place_forget()
        self.docname.place_forget()
        self.save_button.place_forget()
        self.document = Document()
        self.document.styles['Normal'].font.name = 'Times New Roman'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # self.docTk.destroy()



    def getDocxThread(self):
        t1 = threading.Thread(target=self.getDocx)
        t1.setDaemon(True)
        t1.start()

    def getDocx(self):
        vaild = self.vaild_bible()
        if (vaild):
            day = self.this_day_var.get()
            self.data_item[day] = {}
            self.data_item[day]["课程内容"] = self.kcnr_var.get()
            self.data_item[day]["叙述经文"] = self.xsjw.get(0.0, END).strip()
            self.data_item[day]["读经进度"] = self.djjd_var.get()


            self.handle_bible()

    def renderData(self,flag):
        json = {"周一": 0, "周二": 1, "周三": 2, "周四": 3, "周五": 4, "周六": 5, "主日": 6}
        day = self.this_day_var.get()
        self.data_item[day] = {}
        self.data_item[day]["课程内容"]=self.kcnr_var.get()
        self.data_item[day]["叙述经文"] = self.xsjw.get(0.0,END).strip()
        self.data_item[day]["读经进度"] = self.djjd_var.get()

        days = ["周一", "周二", "周三", "周四", "周五", "周六", "主日"]

        for item in days:
            if (item == day):
                self.this_day_var.set(days[days.index(item) + (1 if flag else -1)])

        current_day = self.this_day_var.get()

        if(current_day in self.data_item):
            if (len(self.search_list)>0):
                self.kcnr_var.set(self.search_list[json[current_day]])
            else:
                self.kcnr_var.set(self.data_item[current_day]["课程内容"])
            self.xsjw.delete(0.0, END)
            self.xsjw.insert(END,self.data_item[current_day]["叙述经文"])
            self.djjd_var.set(self.data_item[current_day]["读经进度"])
        else:
            if (len(self.search_list)>0):
                self.kcnr_var.set(self.search_list[json[current_day]])
            else:
                self.kcnr_var.set("")
            self.xsjw.delete(0.0, END)
            self.djjd_var.set("无")
        self.xsjw.focus()
        if(current_day=="周一"):
            self.button1.place_forget()
            self.button2.place(x=360, y=400)
            self.button3.place_forget()
            self.top.unbind('<KeyPress-Up>')
            self.top.bind('<KeyPress-Down>', self.KeyPressDown)
            self.top.unbind('<Return>')
        elif(current_day=="主日"):
            self.button1.place(x=200, y=400)
            self.button3.place(x=360, y=400)
            self.button2.place_forget()
            self.top.bind('<KeyPress-Up>', self.KeyPressUp)
            self.top.unbind('<KeyPress-Down>')
            self.top.bind('<Return>', self.KeyPressReturn)
        else:
            self.button1.place(x=200, y=400)
            self.button2.place(x=360, y=400)
            self.button3.place_forget()
            self.top.bind('<KeyPress-Up>', self.KeyPressUp)
            self.top.bind('<KeyPress-Down>', self.KeyPressDown)
            self.top.unbind('<Return>')


    def prev(self):

        self.renderData(False)

    # 进度条的显示格式
    def progress_bar1(self, now_schedule, all_schedule):
        self.canvas.coords(self.fill_rec, (5, 5, 6 + (now_schedule / all_schedule) * 150, 35))
        self.docTk.update()
        self.zxjd.config(text=str(round(now_schedule / all_schedule * 100, 2)) + '%')
        if round(now_schedule / all_schedule * 100, 2) == 100.00:
            self.zxjd.config(text="完成")


    def next(self):
        vaild = self.vaild_bible()
        if(vaild):
            self.renderData(True)
        # self.renderData(True)

if __name__ == '__main__':
    JwfwClient()