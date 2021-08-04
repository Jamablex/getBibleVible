import json

from docx import Document
import re

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

bible_data = Document("./data/bible.docx")

result_dict = {}
verse_dict = {}

verse_num = 0
node_num = 0
current_volume = "创世记"
current_verse = ""

index_dict = {"1": "创世记", "2": "出埃及记", "3": "利未记", "4": "民数记", "5": "申命记", "6": "约书亚记", "7": "士师记", "8": "路得记",
              "9": "撒母耳记上", "10": "撒母耳记下", "11": "列王记上",
              "12": "列王记下", "13": "历代志上", "14": "历代志下", "15": "以斯拉记", "16": "尼希米记", "17": "以斯帖记", "18": "约伯记",
              "19": "诗篇",
              "20": "箴言", "21": "传道书", "22": "雅歌", "23": "以赛亚书",
              "24": "耶利米书", "25": "耶利米哀歌", "26": "以西结书", "27": "但以理书", "28": "何西阿书", "29": "约珥书", "30": "阿摩司书",
              "31": "俄巴底亚书",
              "32": "约拿书", "33": "弥迦书", "34": "那鸿书", "35": "哈巴谷书",
              "36": "西番雅书", "37": "哈该书", "38": "撒迦利亚书", "39": "玛拉基书", "40": "马太福音", "41": "马可福音", "42": "路加福音",
              "43": "约翰福音",
              "44": "使徒行传", "45": "罗马书", "46": "哥林多前书", "47": "哥林多后书",
              "48": "加拉太书", "49": "以弗所书", "50": "腓立比书", "51": "歌罗西书", "52": "帖撒尼罗迦前书", "53": "帖撒尼罗迦后书", "54": "提摩太前书",
              "55": "提摩太后书", "56": "提多书", "57": "腓利门书", "58": "希伯来书",
              "59": "雅各书", "60": "彼得前书", "61": "彼得后书", "62": "约翰一书", "63": "约翰二书", "64": "约翰三书", "65": "犹大书",
              "66": "启示录"}

row = 0

for paragraph in bible_data.paragraphs:
    for key, value in index_dict.items():
        v_re = r"" + index_dict[key] + "第.*[章篇]"
        match = re.match(v_re, paragraph.text.strip(), re.I)
        if (match is not None):
            split_flag = ""
            if (index_dict[key] != "诗篇"):
                split_flag = "章"
            else:
                split_flag = "篇"
            verse = paragraph.text.strip().split(index_dict[key] + "第")[1].split(split_flag)[0]
            if (current_volume == index_dict[key]):
                verse_dict[verse] = row
            else:
                result_dict[current_volume] = verse_dict
                verse_dict = {}
                verse_dict[verse] = row
                current_volume = index_dict[key]
    row += 1
    # if(row ==5000):
    #     break

result_dict[current_volume] = verse_dict

result = json.dumps(result_dict)
with open(r'./data/result_dict.json', 'w')as f:
    f.write(result)


dict = {}
v_dict = {}
old_v_dict = {}
old_dict = {}
row = 0
nn = 0
v_row = "一"
flag = True
current_key = "创世记"

current_verse_row = result_dict["创世记"][v_row]
for k, v in result_dict.items():
    for node, verse_row in v.items():
        if current_verse_row == result_dict[k][node]:
            continue
        for paragraph in bible_data.paragraphs:
            row = row + 1
            if current_verse_row < row < result_dict[k][node]:
                node_re = r"[0-9]{1,3}:[0-9]{1,3}.*"
                match = re.match(node_re, paragraph.text.strip(), re.I)
                if (match is not None):
                    nn = nn + 1
            if row >= result_dict[k][node]:
                current_verse_row = result_dict[k][node]
                if (flag):
                    old_v_dict[v_row] = nn
                    flag = False
                else:
                    v_dict[v_row] = nn
                v_row = node
                row = 0
                nn = 0
                break
    if k == "创世记":
        v_dict["一"] = 31
        dict[current_key] = v_dict
        old_dict = v_dict
    elif k == "启示录":
        v_dict["二十二"] = 21
        dict[k] = v_dict
    else:
        for key, value in old_dict.items():
            old_v_dict[key] = old_dict[key]
        old_dict = v_dict
        dict[current_key] = old_v_dict
    old_v_dict = {}
    current_key = k
    v_dict = {}
    flag = True

print(dict)
str = json.dumps(dict)
with open(r'./data/dict.json', 'w')as f:
    f.write(str)
# for paragraph in bible_data.paragraphs:
#     n_re = r"[0-9]{1,3}:[0-9]{1,3}"
