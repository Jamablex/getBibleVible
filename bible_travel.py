from docx import Document
import re

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from cmd_colors import Color


bible_data = Document("./data/bible_travel.docx")

document = Document()

def p_format(p):
    paragraph_format = p.paragraph_format
    paragraph_format.space_before = Pt(0)  # 上行间距
    paragraph_format.space_after = Pt(0)  # 下行间距
    paragraph_format.line_spacing = Pt(0)  # 行距

for paragraph in bible_data.paragraphs:
    travel_re = r'(周|主)[一二三四五六日]　.*|第[一二三四五六七八九十]课　.*|第十[一二三四五六七八]课　.*'

    match = re.match(travel_re, paragraph.text.strip(), re.I)
    if match is not None:
        p = document.add_paragraph(paragraph.text)
        p_format(p)

document.save("bible_travel_handle.docx")

