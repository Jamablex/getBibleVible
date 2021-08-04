import json
import os
import time
import tkinter
from tkinter import *


from docx import Document
import re

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from cmd_colors import Color
from progressbar import *





# 加载文档数据
bible_data = Document("../data/bible.docx")
bible_travel_data = Document("../data/bible_travel_handle.docx")
filename = r'../data/dict.json'
fn = r'../data/result_dict.json'

# 验证章节是否有效的dict
valid_dict = json.load(open(filename))
row_dict = json.load(open(fn))


# 验证是否有效
def valid(bible_volume_name, bible_verse_node_str, crt_volume, crt_verse):
    is_valid = True
    # 卷名+章+节格式
    if bible_volume_name != "":
        crt_volume = bible_volume_name
        # 输入卷名不在bible_dict中代表输入错误
        if bible_volume_name not in bible_dict.keys():
            is_valid = False
            return is_valid, crt_volume, crt_verse
        split_flag = ""
        if bible_volume_name == "诗篇":
            split_flag = "篇"
        else:
            split_flag = "章"
        # 分割章节
        verse = bible_verse_node_str.split(split_flag)
        if len(verse) == 2:
            crt_verse = verse[0]
        # 分割后章节都存在的情况
        if len(verse) == 2 and verse[1] != "" and verse[1][-1] != "" and verse[1][-1] != "。":
            if verse[0].isdigit() is False or verse[1].split("节")[0].isdigit() is False:
                is_valid = False
                return is_valid, crt_volume, crt_verse
            if int(verse[0]) > len(valid_dict[crt_volume]) or int(verse[0]) <= 0:
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
            crt_verse = verse[0]
            if (int(verse[1].split("节")[0]) > valid_dict[crt_volume][dealVerseName(verse[0])] or int(
                verse[1].split("节")[0]) <= 0):
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
    # 章+节  章   节 等情况
    else:
        split_flag = ""
        if crt_volume == "":
            is_valid = False
            return is_valid, crt_volume
        if crt_volume == "诗篇":
            split_flag = "篇"
        else:
            split_flag = "章"
        verse = bible_verse_node_str.split(split_flag)
        if len(verse) == 2:
            crt_verse = verse[0]
        # 章+节情况
        if len(verse) == 2 and verse[1] != "" and verse[1][-1] != "" and verse[1][-1] != "。":
            if verse[0].isdigit() is False or verse[1].split("节")[0].isdigit() is False:
                is_valid = False
                return is_valid, crt_volume, crt_verse
            if verse[0] != "" and int(verse[0]) > len(valid_dict[crt_volume]) or int(verse[0]) <= 0:
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
            crt_verse = verse[0]
            if (int(verse[1].split("节")[0]) > valid_dict[crt_volume][dealVerseName(verse[0])] or int(
                verse[1].split("节")[0]) <= 0):
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
        # 章
        elif len(verse) == 2 and (verse[1] == "" or verse[1] == "。"):
            if verse[0].isdigit() is False:
                is_valid = False
                return is_valid, crt_volume, crt_verse
            if int(verse[0]) > len(valid_dict[crt_volume]) or int(verse[0]) <= 0:
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的章数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
            crt_verse = verse[0]
        # 节
        else:
            if verse[0].split("节")[0].isdigit() is False:
                is_valid = False
                return is_valid, crt_volume, crt_verse
            if (int(verse[0].split("节")[0]) > valid_dict[crt_volume][dealVerseName(crt_verse)] or int(
                verse[0].split("节")[0]) <= 0):
                clr.print_yellow_text("检测到“" + crt_volume + bible_verse_node_str + "”的节数输入有误，请重新输入！", False)
                is_valid = "2"
                return is_valid, crt_volume, crt_verse
    return is_valid, crt_volume, crt_verse


def bible_valid(bible_scope):
    crt_volume = ""
    crt_verse = ""
    is_valid = True
    if "。" in bible_scope:
        bible_scope = bible_scope[0:-1]
    scopes = re.split("[,]|[，]", bible_scope)
    if bible_scope != "无":
        for scope in scopes:
            before = ""
            after = ""
            if "~" in scope and len(scope.split("~")) == 2:
                before, after = scope.split("~")
            elif "～" in scope and len(scope.split("～")) == 2:
                before, after = scope.split("～")
            else:
                before = scope
                after = scope
            bible_is_valid, new_crt_volume, new_crt_verse = valid(dealScope(before)[0], dealScope(before)[1],
                                                                  crt_volume, crt_verse)
            crt_volume = new_crt_volume
            crt_verse = new_crt_verse
            if bible_is_valid is False:
                is_valid = False
                break
            if bible_is_valid == "2":
                is_valid = "1"
                break
            bible_is_valid, new_crt_volume1, new_crt_verse1 = valid(dealScope(after)[0], dealScope(after)[1],
                                                                    crt_volume, crt_verse)
            crt_volume = new_crt_volume1
            crt_verse = new_crt_verse1
            if bible_is_valid is False:
                is_valid = False
                break

            if bible_is_valid == "2":
                is_valid = "1"
                break
    else:
        is_valid = "0"
    return is_valid


# 删除文档段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


# 将输入分割为卷名和章节
def dealScope(scope):
    bible_volume_name = ""
    bible_verse_node_str = ""
    is_number = False
    for idx in range(0, len(scope)):
        for n in range(0, len(number_list)):
            if scope[idx] == number_list[n]:
                is_number = True
                break
        if is_number is False:
            bible_volume_name += scope[idx]
        else:
            bible_verse_node_str += scope[idx]
    return bible_volume_name, bible_verse_node_str


# 判断是否为节的情况
def node_is_none(scope):
    is_node = False
    if len(scope) == 1:
        if scope[-1] == "节":
            is_node = True
    else:
        if scope[-1] == "节" or scope[-2] == "节":
            is_node = True
    for n in range(0, len(number_list)):
        if scope[-1] == number_list[n]:
            is_node = True
            break

    return is_node


# 处理章数
def dealVerseName(verse_name):
    verse_handle_name = ""
    if len(verse_name) == 1:
        verse_handle_name = num_dict[verse_name]
    elif len(verse_name) == 2:
        if verse_name[0] == "1":
            verse_handle_name = "十" + num_dict[verse_name[1]]
        else:
            verse_handle_name = num_dict[verse_name[0]] + "十" + num_dict[verse_name[1]]
    else:
        if verse_name[1] == "0" and verse_name[2] == "0":
            verse_handle_name = "一○○"
        elif verse_name[1] == "0":
            verse_handle_name = num_dict[verse_name[0]] + "○" + num_dict[verse_name[2]]
        elif verse_name[2] == "0":
            verse_handle_name = num_dict[verse_name[0]] + num_dict[verse_name[1]] + "○"
        else:
            verse_handle_name = num_dict[verse_name[0]] + num_dict[verse_name[1]] + num_dict[verse_name[2]]
    return verse_handle_name


# dict根据val反取key
def get_key(d, value):
    return [k for k, v in d.items() if v == value]


# 处理写入的文章段落
def p_format(p):
    paragraph_format = p.paragraph_format
    paragraph_format.space_before = Pt(0)  # 上行间距
    paragraph_format.space_after = Pt(0)  # 下行间距
    paragraph_format.line_spacing = Pt(0)  # 行距


# 根据输入课程内容检索每周课程
def w_search(c_input):
    input_split = re.split("[ ]|[　]", c_input)
    c_input = input_split[0] + "　" + input_split[1]
    search_list = []
    index = 0
    for paragraph in bible_travel_data.paragraphs:
        index = index + 1
        if paragraph.text.strip() == c_input:
            for idx in range(0, 7):
                search_list.append(bible_travel_data.paragraphs[index + idx].text.strip())
    return search_list


# cmd颜色改变类
clr = Color()

# 卷名以及简写
bible_dict = {"创世记": "创", "出埃及记": "出", "利未记": "利", "民数记": "民", "申命记": "申", "约书亚记": "书", "士师记": "士", "路得记": "得",
              "撒母耳记上": "撒上", "撒母耳记下": "撒下", "列王记上": "王上",
              "列王记下": "王下", "历代志上": "代上", "历代志下": "代下", "以斯拉记": "拉", "尼希米记": "尼", "以斯帖记": "斯", "约伯记": "伯", "诗篇": "诗",
              "箴言": "箴", "传道书": "传", "雅歌": "歌", "以赛亚书": "赛",
              "耶利米书": "耶", "耶利米哀歌": "哀", "以西结书": "结", "但以理书": "但", "何西阿书": "何", "约珥书": "珥", "阿摩司书": "摩", "俄巴底亚书": "俄",
              "约拿书": "拿", "弥迦书": "弥", "那鸿书": "鸿", "哈巴谷书": "哈",
              "西番雅书": "番", "哈该书": "该", "撒迦利亚书": "亚", "玛拉基书": "玛", "马太福音": "太", "马可福音": "可", "路加福音": "路", "约翰福音": "约",
              "使徒行传": "徒", "罗马书": "罗", "哥林多前书": "林前", "哥林多后书": "林后",
              "加拉太书": "加", "以弗所书": "弗", "腓立比书": "腓", "歌罗西书": "西", "帖撒尼罗迦前书": "帖前", "帖撒尼罗迦后书": "帖后", "提摩太前书": "提前",
              "提摩太后书": "提后", "提多书": "多", "腓利门书": "门", "希伯来书": "来",
              "雅各书": "雅", "彼得前书": "彼前", "彼得后书": "彼后", "约翰一书": "约壹", "约翰二书": "约贰", "约翰三书": "约叁", "犹大书": "犹", "启示录": "启"}

index_dict = {"1": "创世记", "2": "出埃及记", "3": "利未记", "4": "民数记", "5": "申命记", "6": "约书亚记", "7": "士师记", "8": "路得记",
              "9": "撒母耳记上", "10": "撒母耳记下", "11": "列王记上",
              "12": "列王记下", "13": "历代志上", "14": "历代志下", "15": "以斯拉记", "16": "尼希米记", "17": "以斯帖记", "18": "约伯记",
              "19": "诗篇",
              "20": "箴言", "21": "传道书", "22": "雅歌", "23": "以赛亚书",
              "24": "耶利米书", "25": "耶利米哀歌", "26": "以西结书", "27": "但以理书", "28": "何西阿书", "29": "约珥书", "30": "阿摩司书",
              "31": "俄巴底亚书",
              "32": "约拿书", "33": "弥迦书", "34": "那鸿书", "35": "哈巴谷书",
              "36": "西番雅书", "37": "哈该书", "38": "撒迦利亚书", "39": "玛拉基书", "40": "马太福音", "41": "马可福音", "42": "路加福音",
              "43": "约翰福音",
              "44": "使徒行传", "45": "罗马书", "46": "哥林多前书", "47": "哥林多后书",
              "48": "加拉太书", "49": "以弗所书", "50": "腓立比书", "51": "歌罗西书", "52": "帖撒尼罗迦前书", "53": "帖撒尼罗迦后书", "54": "提摩太前书",
              "55": "提摩太后书", "56": "提多书", "57": "腓利门书", "58": "希伯来书",
              "59": "雅各书", "60": "彼得前书", "61": "彼得后书", "62": "约翰一书", "63": "约翰二书", "64": "约翰三书", "65": "犹大书",
              "66": "启示录"}

number_list = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]

num_dict = {"0": "", "1": "一", "2": "二", "3": "三", "4": "四", "5": "五", "6": "六", "7": "七", "8": "八", "9": "九",
            "10": "十"}
# paragraph = document.add_paragraph(bible_dict[bible_name])
current_volume_name = ""
current_verse = ""
canel = False
is_change = False
is_only_verse = False

while True:
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # try:
    clr.print_red_text("请输入课程内容，例：“第六课　大卫受试验以被成全”", False)
    c_input = clr.print_blue_text("输入：", True)
    if c_input == "":
        clr.print_red_text("检测到课程内容输入为空，请重新输入！", False)
        continue
    if len(re.split("[ ]|[　]", c_input)) == 1:
        clr.print_red_text("检测到课程内容输入有误，请重新输入！", False)
        continue
    search_list = w_search(c_input)
    row_one = document.add_paragraph(c_input)
    num = 0
    # num用来判断是第几周的课程
    while num < 7:
        if canel != is_change:
            num = num - 1
            is_change = False
        if len(search_list) == 0:
            if num != 6:
                clr.print_red_text(
                    "请输入周" + num_dict[str(num + 1)] + "的课程内容，例：“周" + num_dict[str(num + 1)] + "　神主宰一切环境为着成全大卫”",
                    False)
            else:
                clr.print_red_text("请输入主日的课程内容，例：“主日  神主宰一切环境为着成全大卫”", False)
        else:
            clr.print_red_text("检测到当前课程内容应为“" + search_list[num] + "”，无需更改输入回车即可，或输入修改后的课程内容", False)
        num = num + 1
        w_input = clr.print_blue_text("输入：", True)
        row_two = ""
        if w_input == "":
            row_two = document.add_paragraph(search_list[num - 1])
        else:
            row_two = document.add_paragraph(w_input)
        scope = ""
        scope_is_none = False
        while True:
            clr.print_red_text("请输入叙述经文范围,输入格式参考Readme.md文件，例：“创世记1章10节~3章5节”，多个叙述经文范围请按照“逗号”分割", False)
            scope = clr.print_blue_text("输入：", True).strip()
            if scope == "":
                clr.print_red_text("检测到叙述经文输入为空，请重新输入！", False)
                continue
            scope_is_valid = bible_valid(scope)
            if scope_is_valid is False:
                clr.print_red_text("检测到叙述经文输入有误，请重新输入！", False)
                continue
            elif scope_is_valid == "0":
                scope_is_none = True
                print("正在执行，请稍等......")
                row_three = document.add_paragraph()
                title = row_three.add_run("叙述经文：")
                title.font.color.rgb = RGBColor(255, 0, 0)
                title.bold = True
                content = row_three.add_run(scope)
                content.font.color.rgb = RGBColor(255, 0, 0)
                content.font.highlight_color = MSO_THEME_COLOR_INDEX.ACCENT_3
                content.bold = True
                p_format(row_one)
                p_format(row_two)
                p_format(row_three)
            elif scope_is_valid == "1":
                continue
            break
        if scope_is_none:
            continue
        try:
            row_three = document.add_paragraph()
            title = row_three.add_run("叙述经文：")
            title.font.color.rgb = RGBColor(255, 0, 0)
            title.bold = True
            content = row_three.add_run(scope)
            content.font.color.rgb = RGBColor(255, 0, 0)
            content.font.highlight_color = MSO_THEME_COLOR_INDEX.ACCENT_3
            content.bold = True
            p_format(row_one)
            p_format(row_two)
            p_format(row_three)
            if "。" in scope:
                scope = scope[0:-1]
            bible_read_list = re.split("[,]|[，]", scope)
            widgets = ['任务进度: ', Percentage(), ' ', Bar(marker='#', left='[', right=']', fill='-'), ' ',
                       ' ', ETA(), ' ']
            pbar = ProgressBar(widgets=widgets).start()
            pbar.update(0 / len(bible_read_list) * 100)
            for i in range(0, len(bible_read_list)):
                pbar.update(int(((i + 1 / 10) / len(bible_read_list) * 100)))
                bible_before_volume_name = ""
                bible_before_verse_name = ""
                bible_before_node_name = ""
                bible_after_volume_name = ""
                bible_after_verse_name = ""
                bible_after_node_name = ""
                if "~" in bible_read_list[i]:
                    scope_before = bible_read_list[i].split("~")
                else:
                    scope_before = bible_read_list[i].split("～")
                # 有范围  后与前各取，无范围，后与前相同
                if len(scope_before) == 2:
                    bible_volume_before, bible_verse_node_before = dealScope(scope_before[0])
                    bible_volume_after, bible_verse_node_after = dealScope(scope_before[1])
                else:
                    bible_volume_before, bible_verse_node_before = dealScope(scope_before[0])
                    bible_volume_after, bible_verse_node_after = dealScope(scope_before[0])
                if len(scope_before) == 2:
                    # 有章节且卷名不为空的情况
                    if bible_volume_before != "":
                        bible_before_volume_name = bible_volume_before
                        current_volume_name = bible_before_volume_name
                    else:
                        bible_before_volume_name = current_volume_name
                    split_flag = ""
                    if bible_before_volume_name != "诗篇":
                        split_flag = "章"
                    else:
                        split_flag = "篇"
                    # 卷名+节情况
                    if (len(bible_verse_node_before.split(split_flag)) == 2 and
                        bible_verse_node_before.split(split_flag)[1] != "" and
                        bible_verse_node_before.split(split_flag)[1] != "。"):
                        bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                        current_verse = bible_before_verse_name
                        bible_before_node_name = bible_verse_node_before.split(split_flag)[1].split("节")[0]
                    else:
                        is_node = node_is_none(bible_verse_node_before)
                        after_is_node = node_is_none(bible_verse_node_after)
                        # 只有节
                        if is_node and after_is_node:
                            bible_before_verse_name = current_verse
                            bible_before_node_name = bible_verse_node_before.split("节")[0]
                        # 只有章
                        else:
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            bible_before_node_name = "1"
                    # 与前对应
                    if bible_volume_after != "":
                        bible_after_volume_name = bible_volume_after
                        # 章+节
                        if (len(bible_verse_node_after.split(split_flag)) == 2 and
                            bible_verse_node_after.split(split_flag)[1] != "" and
                            bible_verse_node_after.split(split_flag)[1] != "。"):
                            bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                            bible_after_node_name = bible_verse_node_after.split(split_flag)[1].split("节")[0]
                        else:
                            is_node = node_is_none(bible_verse_node_after)
                            # 只有节
                            if is_node:
                                bible_after_verse_name = current_verse
                                bible_after_node_name = bible_verse_node_after.split("节")[0]
                            # 只有章
                            else:
                                bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                                bible_after_node_name = "999"
                    # 无卷名情况
                    else:
                        bible_after_volume_name = current_volume_name
                        # 章+节
                        if (len(bible_verse_node_after.split(split_flag)) == 2 and
                            bible_verse_node_after.split(split_flag)[1] != "" and
                            bible_verse_node_after.split(split_flag)[1] != "。"):
                            bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                            current_verse = bible_after_verse_name
                            bible_after_node_name = bible_verse_node_after.split(split_flag)[1].split("节")[0]
                        else:
                            is_node = node_is_none(bible_verse_node_after)
                            # 节
                            if is_node:
                                bible_after_verse_name = current_verse
                                bible_after_node_name = bible_verse_node_after.split("节")[0]
                            # 章
                            else:
                                bible_after_verse_name = bible_verse_node_after.split(split_flag)[0]
                                bible_after_node_name = "999"
                # 后与前相同
                else:
                    # 卷名+章节
                    if bible_volume_before != "":
                        if bible_volume_before != "诗篇":
                            split_flag = "章"
                        else:
                            split_flag = "篇"
                        if (len(bible_verse_node_before.split(split_flag)) == 2 and
                            bible_verse_node_before.split(split_flag)[1] != "" and
                            bible_verse_node_before.split(split_flag)[1] != "。"):
                            bible_before_volume_name = bible_volume_before
                            current_volume_name = bible_before_volume_name
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            current_verse = bible_before_verse_name
                            bible_before_node_name = bible_verse_node_before.split(split_flag)[1].split("节")[0]
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = bible_before_node_name
                        else:
                            bible_before_volume_name = bible_volume_before
                            current_volume_name = bible_before_volume_name
                            bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                            current_verse = bible_before_verse_name
                            bible_before_node_name = "1"
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = "999"
                    # 章节
                    else:
                        bible_before_volume_name = current_volume_name
                        if bible_before_volume_name != "诗篇":
                            split_flag = "章"
                        else:
                            split_flag = "篇"
                        if (len(bible_verse_node_before.split(split_flag)) == 2 and
                            bible_verse_node_before.split(split_flag)[1] != "" and
                            bible_verse_node_before.split(split_flag)[1] != "。"):
                            bible_before_verse_name = scope_before[0].split(split_flag)[0]
                            current_verse = bible_before_verse_name
                            bible_before_node_name = scope_before[0].split(split_flag)[1].split("节")[0]
                            bible_after_volume_name = bible_before_volume_name
                            bible_after_verse_name = bible_before_verse_name
                            bible_after_node_name = bible_before_node_name
                        else:
                            is_node = node_is_none(bible_verse_node_before)
                            # 节
                            if is_node:
                                bible_before_verse_name = current_verse
                                bible_before_node_name = bible_verse_node_before.split("节")[0]
                                bible_after_volume_name = bible_before_volume_name
                                bible_after_verse_name = bible_before_verse_name
                                bible_after_node_name = bible_before_node_name
                            # 章
                            else:
                                bible_before_verse_name = bible_verse_node_before.split(split_flag)[0]
                                current_verse = bible_before_verse_name
                                bible_before_node_name = "1"
                                bible_after_volume_name = bible_before_volume_name
                                bible_after_verse_name = bible_before_verse_name
                                bible_after_node_name = "999"
                before_verse_handle_name = ""
                after_verse_handle_name = ""
                try:
                    # 章数格式化
                    before_verse_handle_name = dealVerseName(bible_before_verse_name)
                    after_verse_handle_name = dealVerseName(str(int(bible_after_verse_name) + 1))
                except:
                    clr.print_red_text("检测到叙述经文范围可能输入错误，请确认是否重新输入！重新输入：1  或按任意键继续执行！", False)
                    operate = clr.print_blue_text("输入：", True).strip()
                    if operate == "1":
                        delete_paragraph(document.paragraphs[len(document.paragraphs) - 1])
                        delete_paragraph(document.paragraphs[len(document.paragraphs) - 2])
                        is_change = True
                        break
                    else:
                        break
                if bible_before_volume_name != "诗篇":
                    bible_before_volume_verse = bible_before_volume_name + "第" + before_verse_handle_name + "章"
                else:
                    bible_before_volume_verse = bible_before_volume_name + "第" + before_verse_handle_name + "篇"
                if bible_after_volume_name != "诗篇":
                    bible_after_volume_verse = bible_after_volume_name + "第" + after_verse_handle_name + "章"
                else:
                    bible_after_volume_verse = bible_after_volume_name + "第" + after_verse_handle_name + "篇"
                begin_row = 0
                done_row = 0
                is_volume_end = True
                is_bible_end = True
                # 先计算读经范围是哪一章到哪一章
                pbar.update(int(((i + 2 / 10) / len(bible_read_list) * 100)))
                before_bar_count = 0
                for paragraph in bible_data.paragraphs:
                    begin_row = begin_row + 1
                    if paragraph.text.strip() == bible_before_volume_verse:
                        break
                    if row_dict[bible_before_volume_name][before_verse_handle_name] > 5000:
                        before_bar_num = row_dict[bible_before_volume_name][before_verse_handle_name] // 5000
                        if begin_row % 5000 == 0:
                            before_bar_count = before_bar_count + 1
                            pbar.update(int(((i + 2 / 10 + (2.5 / 10) * (before_bar_count / before_bar_num)) / len(
                                bible_read_list) * 100)))
                pbar.update(int(((i + 4.5 / 10) / len(bible_read_list) * 100)))
                after_bar_count = 0
                for paragraph in bible_data.paragraphs:
                    done_row = done_row + 1
                    if paragraph.text.strip() == bible_after_volume_verse:
                        is_volume_end = False
                        is_bible_end = False
                        break
                    after_bar_num = ""
                    if after_verse_handle_name in row_dict[bible_after_volume_name] and \
                        row_dict[bible_after_volume_name][after_verse_handle_name] > 5000:
                        after_bar_num = row_dict[bible_after_volume_name][after_verse_handle_name] // 5000
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            pbar.update(int(((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list) * 100)))
                    elif after_verse_handle_name == "二十三" and bible_after_volume_name == "启示录":
                        after_bar_num = row_dict[bible_after_volume_name]["二十二"] // 5000
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            pbar.update(int(((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list) * 100)))
                    else:
                        key = get_key(index_dict, bible_after_volume_name)[0]
                        after_bar_num = row_dict[index_dict[str(int(key) + 1)]]["一"]
                        if done_row % 5000 == 0:
                            after_bar_count = after_bar_count + 1
                            pbar.update(int(((i + 4.5 / 10 + (2.5 / 10) * (after_bar_count / after_bar_num)) / len(
                                bible_read_list) * 100)))

                pbar.update(int(((i + 7 / 10) / len(bible_read_list) * 100)))
                # 读经范围到启示录的情况
                if is_volume_end:
                    done_row = 0
                    key = get_key(index_dict, bible_after_volume_name)[0]
                    if int(key) + 1 > 66:
                        is_bible_end = True
                    else:
                        bible_after_volume_verse = index_dict[str(int(key) + 1)] + "第一章"
                        for paragraph in bible_data.paragraphs:
                            done_row = done_row + 1
                            if paragraph.text.strip() == bible_after_volume_verse:
                                is_bible_end = False
                                break
                # 如果是启示录，直接预设最大行数
                if is_bible_end:
                    done_row = 49999
                begin_node = begin_row
                end_node_num = begin_row
                end_node = end_node_num
                row = 0
                # 范围进一步精确到哪一章哪一节
                for paragraph in bible_data.paragraphs:
                    row = row + 1
                    if begin_row < row < done_row:
                        begin = bible_before_verse_name + ":" + bible_before_node_name
                        begin_re = r'' + begin + '.*'
                        match = re.match(begin_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            begin_node = begin_node + 1
                            break
                        else:
                            begin_node = begin_node + 1
                pbar.update(int(((i + 8 / 10) / len(bible_read_list) * 100)))
                row = 0
                for paragraph in bible_data.paragraphs:
                    row = row + 1
                    if begin_row < row < done_row:
                        end = bible_after_verse_name + ":" + bible_after_node_name
                        end_re = r'' + end + '.*'
                        match = re.match(end_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            end_node = end_node + 1
                            break
                        else:
                            end_node = end_node + 1
                row = 0
                volume_name = bible_before_volume_name
                pbar.update(int(((i + 9 / 10) / len(bible_read_list) * 100)))
                # 检索中间的段落，并写入新文档中
                for paragraph in bible_data.paragraphs:
                    row = row + 1
                    if begin_node <= row <= end_node:
                        verse_name_re = ""
                        if volume_name != "诗篇":
                            verse_name_re = r'[\u4E00-\u9FA5]+第[\u4E00-\u9FA5]+章'
                        else:
                            verse_name_re = r'[\u4E00-\u9FA5]+第[\u4E00-\u9FA5]+篇'
                        vn = re.match(verse_name_re, paragraph.text.strip(), re.I)
                        if vn is not None:
                            volume_name = paragraph.text.split("第")[0]
                        vb_re = r'[0-9]{1,3}:[0-9]{1,3}.*'
                        match = re.match(vb_re, paragraph.text.strip(), re.I)
                        if match is not None:
                            head = paragraph.text.split("　")[0]
                            content = paragraph.text.split("　")[1]
                            deal_content = re.sub('[0-9]', "", content)
                            label = bible_dict[volume_name] + head
                            row_result = "　" + deal_content
                            p = document.add_paragraph()
                            run = p.add_run(label)
                            run.font.color.rgb = RGBColor(0, 112, 192)
                            run.bold = True
                            p.add_run(row_result)
                            p_format(p)
                # os.system("cls")
                pbar.update(int(((i + 1) / len(bible_read_list) * 100)))
                if i == len(bible_read_list) - 1:
                    pbar.finish()
            clr.print_red_text("请输入需要读经进度,若没有请输入无", False)
            djjd = clr.print_blue_text("输入：", True).strip()
            p = document.add_paragraph()
            p_title = p.add_run("读经进度：")
            p_title.font.color.rgb = RGBColor(255, 0, 0)
            p_title.bold = True
            run = p.add_run(djjd)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
        except:
            if is_change is False:
                clr.print_red_text("检测到叙述经文可能输入错误，请确认是否重新输入！重新输入：1  或按任意键继续执行！", False)
                operate = clr.print_blue_text("输入：", True).strip()
                if operate == "1":
                    delete_paragraph(document.paragraphs[len(document.paragraphs) - 1])
                    delete_paragraph(document.paragraphs[len(document.paragraphs) - 2])
                    is_change = True

    clr.print_red_text("请输入需要保存的文件名称", False)
    docx_name = clr.print_blue_text("输入：", True).strip()
    document.save(docx_name + ".docx")
    clr.print_red_text("文件生成成功！是否继续执行? 继续：1  或输入任意字符退出 ", False)
    operate = clr.print_blue_text("输入：", True).strip()
    if operate != "1":
        break
    # os.system("cls")
    # except:
    #     clr.print_red_text("请检查文件是否关闭，读经范围等是否输入正确！", True)


